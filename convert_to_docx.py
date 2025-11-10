"""
Convert COMPLETE_PROJECT_GUIDE.md to .docx format
Run this script: python convert_to_docx.py
"""

import re
import os

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx library...")
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_markdown_to_docx(md_file, output_file):
    """Convert markdown file to Word document"""
    
    print(f"Reading {md_file}...")
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    print("Creating Word document...")
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    lines = content.split('\n')
    in_code_block = False
    code_lines = []
    skip_next = False
    
    i = 0
    total_lines = len(lines)
    
    while i < total_lines:
        line = lines[i]
        
        if skip_next:
            skip_next = False
            i += 1
            continue
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block - add accumulated code
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_text)
                    font = p.runs[0].font
                    font.name = 'Consolas'
                    font.size = Pt(9)
                    p.paragraph_format.left_indent = Inches(0.5)
                code_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Handle main title (first line)
        if i == 0 or (line.startswith('# ') and not line.startswith('##')):
            text = line.replace('# ', '').strip()
            if text:
                p = doc.add_heading(text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Handle H2 headers (## )
        if line.startswith('## '):
            text = line.replace('## ', '')
            text = re.sub(r'^\d+\.\s*', '', text)  # Remove "1. "
            text = re.sub(r'\s*{#.*}$', '', text)  # Remove {#anchor}
            if text.strip():
                doc.add_heading(text.strip(), level=1)
            i += 1
            continue
        
        # Handle H3 headers (### )
        if line.startswith('### '):
            text = line.replace('### ', '')
            text = re.sub(r'^\d+\.\d+\s*', '', text)  # Remove "1.1 "
            if text.strip():
                doc.add_heading(text.strip(), level=2)
            i += 1
            continue
        
        # Handle H4 headers (#### )
        if line.startswith('#### '):
            text = line.replace('#### ', '').strip()
            if text:
                doc.add_heading(text, level=3)
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() in ['---', '___', '***']:
            doc.add_paragraph()
            i += 1
            continue
        
        # Handle bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Clean markdown
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'"\1"', text)
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
            if text.strip():
                doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # Handle numbered lists
        match = re.match(r'^(\d+)\.\s+(.+)', line.strip())
        if match:
            text = match.group(2)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'"\1"', text)
            if text.strip():
                doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Handle tables (add as plain text for simplicity)
        if '|' in line and line.strip().startswith('|'):
            # Simple table as formatted text
            text = line.strip()
            if '-|-' not in text and text != '||':
                p = doc.add_paragraph(text)
                font = p.runs[0].font
                font.name = 'Consolas'
                font.size = Pt(9)
            i += 1
            continue
        
        # Handle regular paragraphs
        if line.strip():
            text = line.strip()
            # Clean markdown
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'"\1"', text)
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
            
            # Skip certain decorative lines
            if any(char in text for char in ['┌', '│', '└', '├', '─', '▼', '┐', '┘', '┬', '┴']):
                p = doc.add_paragraph(text)
                font = p.runs[0].font
                font.name = 'Consolas'
                font.size = Pt(8)
            elif text:
                doc.add_paragraph(text)
        else:
            # Empty line - add space
            doc.add_paragraph()
        
        i += 1
    
    # Save document
    print(f"Saving to {output_file}...")
    doc.save(output_file)
    file_size_mb = os.path.getsize(output_file) / 1024 / 1024
    print(f"[OK] Successfully created {output_file}")
    print(f"  File size: {file_size_mb:.2f} MB")
    print(f"  Location: {os.path.abspath(output_file)}")
    print(f"\n[SUCCESS] You can now open it in Microsoft Word!")

if __name__ == '__main__':
    input_file = 'docs/COMPLETE_PROJECT_GUIDE.md'
    output_file = 'docs/COMPLETE_PROJECT_GUIDE.docx'
    
    if not os.path.exists(input_file):
        print(f"Error: {input_file} not found!")
        print("Make sure you're running this from the BASE-PAPERS directory")
        exit(1)
    
    print("="*60)
    print("  Markdown to Word Converter")
    print("="*60)
    print()
    
    convert_markdown_to_docx(input_file, output_file)
    
    print("\n" + "="*60)
    print("  CONVERSION COMPLETE!")
    print("="*60)
    print(f"\nNext step: Open '{output_file}' in Microsoft Word")
    print("The document is ready for printing or presentation!")
